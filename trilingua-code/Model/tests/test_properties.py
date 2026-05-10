# -*- coding: utf-8 -*-
"""
Property-based tests for translation-layout-quality components.

Uses Hypothesis for property generation.
Each test is tagged with a comment referencing the design property.
"""

import sys
import os
import re

# Ensure the Model package is importable when running pytest from the repo root.
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from hypothesis import given, settings, assume
from hypothesis import strategies as st

from document_translator_v3 import Context_Buffer, Chunk_Splitter, Glossary_Store


# ---------------------------------------------------------------------------
# Tokenizer approximation
# ---------------------------------------------------------------------------
# Loading the real NLLB tokenizer (facebook/nllb-200-distilled-600M) in tests
# would require downloading ~1 GB of model weights and would be very slow.
# Instead we use a word-count approximation that mirrors the budget check in
# _translate_single():
#
#   1 word ≈ 1.3 tokens  →  400-token budget ≈ 307 words
#
# The truncation algorithm in _translate_single() works word-by-word, so the
# approximation is faithful to the actual implementation.

TOKEN_BUDGET = 400
WORDS_PER_TOKEN = 1.3
WORD_BUDGET = int(TOKEN_BUDGET / WORDS_PER_TOKEN)  # 307


def _approx_token_count(text: str) -> int:
    """
    Approximate the NLLB tokenizer token count for *text*.

    Uses the same 1-word ≈ 1.3-tokens heuristic that the production code
    relies on when the real tokenizer is unavailable.  The +2 accounts for
    the BOS/EOS special tokens that the real tokenizer always adds.
    """
    word_count = len(text.split())
    return round(word_count * WORDS_PER_TOKEN) + 2


def _apply_truncation(hint: str, source: str) -> str:
    """
    Replicate the truncation logic from _translate_single() using the
    word-count approximation instead of the real tokenizer.

    If the combined token count of ``hint + " ||| " + source`` exceeds
    TOKEN_BUDGET, words are dropped from the *front* of the hint one at a
    time until the combined count fits within the budget (or the hint is
    exhausted, in which case only the source is used).

    Returns the final combined string that would be passed to the encoder.
    """
    if not hint:
        return source

    combined = hint + " ||| " + source
    if _approx_token_count(combined) <= TOKEN_BUDGET:
        return combined

    # Truncate hint from the front, word by word (oldest context first).
    hint_words = hint.split()
    while hint_words:
        hint_words.pop(0)  # drop oldest word
        if hint_words:
            combined = " ".join(hint_words) + " ||| " + source
        else:
            combined = source
        if _approx_token_count(combined) <= TOKEN_BUDGET:
            break

    return combined


# ===========================================================================
# Property 1: Chunk boundaries fall at sentence endings
# Feature: translation-layout-quality, Property 1: Chunk boundaries fall at sentence endings
# ===========================================================================

# Words that do NOT end with sentence-ending punctuation.
_plain_word = st.text(
    alphabet=st.characters(
        whitelist_categories=("Ll", "Lu"),  # lowercase + uppercase letters
        blacklist_characters=".!?",
    ),
    min_size=1,
    max_size=12,
)

# A sentence-ending punctuation character.
_sentence_end_char = st.sampled_from([".", "!", "?"])


@st.composite
def text_with_sentence_endings(draw, min_tokens=81, max_tokens=300):
    """
    Generate a string of *n* whitespace-delimited tokens (81 ≤ n ≤ 300)
    where sentence-ending punctuation is placed such that:

    1. At least one boundary falls within the first 80 tokens (so a split
       is actually triggered).
    2. Every boundary position is at least ``min_tokens`` (5) tokens from
       the start of its chunk, so that the ``min_tokens`` merging logic in
       Chunk_Splitter does NOT merge a tiny leading chunk into the next one
       (which would produce a non-punctuation-terminated intermediate chunk).

    Concretely, boundaries are placed at positions that are multiples of a
    step size drawn from [10, 40], ensuring each chunk is large enough to
    avoid the minimum-size merge.
    """
    MIN_CHUNK = 5  # must match Chunk_Splitter default min_tokens

    n = draw(st.integers(min_value=min_tokens, max_value=max_tokens))
    words = draw(st.lists(_plain_word, min_size=n, max_size=n))

    # Ensure every word is non-empty.
    assume(all(w for w in words))

    # Choose a step size that guarantees each chunk is >= MIN_CHUNK tokens.
    # Step is in [MIN_CHUNK, 40] so boundaries are well-spaced.
    step = draw(st.integers(min_value=MIN_CHUNK, max_value=40))

    # Place boundaries at step, 2*step, 3*step, ... as long as the position
    # is within the first 80 tokens (so at least one split is triggered) and
    # there are at least MIN_CHUNK tokens remaining after the boundary.
    boundary_positions = set()
    pos = step - 1  # 0-indexed: boundary after word at index `pos`
    while pos < n - MIN_CHUNK:
        boundary_positions.add(pos)
        if pos < 80:
            # We have at least one boundary in the first 80 tokens — good.
            pass
        pos += step

    # If no boundary landed within the first 80 tokens, force one at
    # position (step - 1) clamped to [MIN_CHUNK - 1, 79].
    if not any(p < 80 for p in boundary_positions):
        forced = min(max(step - 1, MIN_CHUNK - 1), min(79, n - MIN_CHUNK - 1))
        boundary_positions.add(forced)

    # Append a sentence-ending character to each chosen word.
    for pos in sorted(boundary_positions):
        punct = draw(_sentence_end_char)
        words[pos] = words[pos] + punct

    return " ".join(words)


@given(text=text_with_sentence_endings())
@settings(max_examples=100)
def test_chunk_boundaries_fall_at_sentence_endings(text):
    """
    **Validates: Requirements 1.1**

    Property 1: Chunk boundaries fall at sentence endings

    For any block of text with more than 80 tokens, every chunk produced by
    Chunk_Splitter.split() — except possibly the final chunk — must end with
    a sentence-ending punctuation character (`.`, `!`, or `?`).
    """
    # Feature: translation-layout-quality, Property 1: Chunk boundaries fall at sentence endings

    splitter = Chunk_Splitter()
    chunks = splitter.split(text, max_tokens=80, hard_cap=150, min_tokens=5)

    # The input always has > 80 tokens, so we expect at least one chunk.
    assert len(chunks) >= 1

    # Every chunk except the last must end with sentence-ending punctuation.
    for i, chunk in enumerate(chunks[:-1]):
        stripped = chunk.rstrip()
        assert stripped and stripped[-1] in {".", "!", "?"}, (
            f"Chunk {i} (of {len(chunks)}) does not end with sentence-ending "
            f"punctuation.\n"
            f"  Chunk text (last 60 chars): ...{stripped[-60:]!r}\n"
            f"  Full text (first 120 chars): {text[:120]!r}"
        )


# ===========================================================================
# Property 2: Chunk content round-trip
# Feature: translation-layout-quality, Property 2: Chunk content round-trip
# ===========================================================================

@given(st.text(alphabet=st.characters(
    # Exclude surrogate characters (Cs), all Unicode control characters (Cc),
    # and all Unicode space/separator categories (Zs, Zl, Zp) EXCEPT the
    # regular ASCII space U+0020.  This ensures the only whitespace character
    # in generated strings is the plain space, so Python's str.split() and
    # Chunk_Splitter's internal split() agree on token boundaries and the
    # fast-path returning [text] verbatim produces the same normalized form.
    blacklist_categories=("Cs", "Cc", "Zs", "Zl", "Zp"),
    whitelist_characters=" ",  # re-allow regular ASCII space
)))
@settings(max_examples=200)
def test_chunk_content_round_trip(text: str):
    """
    **Validates: Requirements 1.6**

    Property 2: Chunk content round-trip

    For any text string, joining all chunks produced by Chunk_Splitter.split()
    with a single space must reproduce the original text after normalising
    runs of whitespace to a single space.

    That is:
        " ".join(splitter.split(text))  ==  " ".join(text.split())

    The generator restricts whitespace to plain ASCII space (U+0020) only,
    matching real PDF/DOCX block text where the only whitespace is spaces.
    Surrogate, control, and Unicode space/separator characters are excluded
    because Python's str.split() treats them as whitespace but they never
    appear in real document blocks.
    """
    # Feature: translation-layout-quality, Property 2: Chunk content round-trip

    # Skip whitespace-only strings — they are never passed to Chunk_Splitter
    # in production (all readers filter blocks to >= 3 meaningful words).
    assume(text.split())  # at least one non-whitespace token

    # Skip strings with leading/trailing whitespace — all document readers
    # strip blocks before passing them to Chunk_Splitter, so the fast-path
    # returning [text] verbatim is correct for stripped inputs only.
    assume(text == text.strip())

    splitter = Chunk_Splitter()
    chunks = splitter.split(text)
    reconstructed = " ".join(chunks)
    expected = " ".join(text.split())
    assert reconstructed == expected, (
        f"Round-trip failed.\n"
        f"  Input:         {text!r}\n"
        f"  Chunks:        {chunks!r}\n"
        f"  Reconstructed: {reconstructed!r}\n"
        f"  Expected:      {expected!r}"
    )


# ===========================================================================
# Property 4: Context hint token budget
# Feature: translation-layout-quality, Property 4: Context hint token budget
# ===========================================================================

# Strategies for generating realistic hint and source strings.
# We use printable ASCII words (no newlines) to keep the word-split
# approximation clean.  Lengths are chosen to exercise both the
# "fits without truncation" and "requires truncation" paths.

_word = st.text(
    alphabet=st.characters(whitelist_categories=("Ll", "Lu", "Nd"),
                           whitelist_characters="-'"),
    min_size=1,
    max_size=15,
).filter(lambda w: w.strip() != "")

_sentence = st.lists(_word, min_size=1, max_size=50).map(" ".join)

_hint_strategy = st.lists(_sentence, min_size=0, max_size=10).map(
    lambda sentences: " ".join(sentences)
)

_source_strategy = st.lists(_sentence, min_size=1, max_size=10).map(
    lambda sentences: " ".join(sentences)
)


@settings(max_examples=200)
@given(hint=_hint_strategy, source=_source_strategy)
def test_context_hint_token_budget(hint: str, source: str):
    """
    **Validates: Requirements 2.3**

    Property 4: Context hint token budget

    For any context hint string and source block string, the combined token
    count of the encoded input passed to the NLLB model shall not exceed 400
    tokens after the truncation logic in _translate_single() is applied.

    The test replicates the word-by-word front-truncation algorithm from
    _translate_single() using a word-count approximation
    (1 word ≈ 1.3 tokens, budget = 400 tokens ≈ 307 words) so that the
    property can be verified quickly without loading the full NLLB tokenizer.
    """
    # Feature: translation-layout-quality, Property 4: Context hint token budget

    # Apply the same truncation logic used in _translate_single().
    result = _apply_truncation(hint, source)

    # The approximated token count of the result must be within budget.
    token_count = _approx_token_count(result)
    assert token_count <= TOKEN_BUDGET, (
        f"Token count {token_count} exceeds budget {TOKEN_BUDGET}.\n"
        f"  hint ({len(hint.split())} words): {hint[:80]!r}...\n"
        f"  source ({len(source.split())} words): {source[:80]!r}...\n"
        f"  result ({len(result.split())} words): {result[:80]!r}..."
    )


@settings(max_examples=100)
@given(hint=_hint_strategy, source=_source_strategy)
def test_context_hint_source_always_preserved(hint: str, source: str):
    """
    **Validates: Requirements 2.3**

    Corollary of Property 4: The source block text is NEVER truncated —
    only the hint is truncated.  After applying the truncation logic, the
    result must end with the original source text.
    """
    # Feature: translation-layout-quality, Property 4: Context hint token budget

    result = _apply_truncation(hint, source)

    # The result must always end with the source text (possibly preceded by
    # a truncated hint and the " ||| " delimiter, or be the source alone).
    assert result.endswith(source), (
        f"Source text was modified or dropped during truncation.\n"
        f"  source: {source!r}\n"
        f"  result: {result!r}"
    )


@settings(max_examples=100)
@given(
    hint=st.lists(_word, min_size=500, max_size=600).map(" ".join),
    source=st.lists(_word, min_size=1, max_size=20).map(" ".join),
)
def test_context_hint_truncation_triggered_for_long_hints(hint: str, source: str):
    """
    **Validates: Requirements 2.3**

    When the hint is very long (500–600 words), the truncation logic MUST
    reduce the combined token count to within the 400-token budget.
    """
    # Feature: translation-layout-quality, Property 4: Context hint token budget

    # Verify the pre-truncation count actually exceeds the budget so we know
    # truncation is genuinely exercised.
    combined_before = hint + " ||| " + source
    assume(_approx_token_count(combined_before) > TOKEN_BUDGET)

    result = _apply_truncation(hint, source)
    token_count = _approx_token_count(result)

    assert token_count <= TOKEN_BUDGET, (
        f"Truncation failed: token count {token_count} still exceeds "
        f"budget {TOKEN_BUDGET} after truncation.\n"
        f"  hint words: {len(hint.split())}\n"
        f"  source words: {len(source.split())}\n"
        f"  result words: {len(result.split())}"
    )


# ===========================================================================
# Property 6: Glossary capitalisation preservation
# Feature: translation-layout-quality, Property 6: Glossary capitalisation preservation
# ===========================================================================

# Strategy: generate a random alphabetic string (no digits, no punctuation)
# to use as a glossary target term.
# Restricted to ASCII letters (a-z, A-Z) so that str.upper() / str.lower() /
# str.title() / str.isupper() / str.islower() / str.istitle() all behave
# predictably.  Unicode letters outside ASCII (e.g. ꟺ U+A7FA) may have no
# case mapping, causing isupper() to return False even after .upper().
_ASCII_LETTERS = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ"
_alpha_word = st.text(
    alphabet=_ASCII_LETTERS,
    min_size=1,
    max_size=20,
)

# Strategy: generate a token in one of the three canonical capitalisation
# patterns — all-caps, title-case, or lowercase.
_cap_pattern = st.sampled_from(["upper", "title", "lower"])


@st.composite
def token_and_target(draw):
    """
    Draw a (token, target_term) pair where:
    - `token` is a non-empty alphabetic string in one of the three canonical
      capitalisation patterns (all-caps, title-case, lowercase).
    - `target_term` is a random non-empty alphabetic string (stored as-is).
    """
    base = draw(_alpha_word)
    assume(base.strip())  # ensure non-empty after stripping

    pattern = draw(_cap_pattern)
    if pattern == "upper":
        token = base.upper()
        # Ensure the token is genuinely all-caps (has at least one cased char).
        assume(token.isupper())
    elif pattern == "title":
        # For title-case to be unambiguous, the base must have at least 2
        # characters so that token.title() has both an uppercase first char
        # and at least one lowercase char.  A single-char title-case token
        # (e.g. "A") satisfies both isupper() and istitle(), which would
        # cause _match_case to take the isupper() branch instead.
        assume(len(base) >= 2)
        token = base.lower()  # start from all-lowercase so title() is clean
        token = token.title()
        # Verify the token is genuinely title-case and NOT all-caps.
        assume(token.istitle() and not token.isupper())
    else:
        token = base.lower()
        # Ensure the token is genuinely all-lowercase (has at least one cased char).
        assume(token.islower())

    target = draw(_alpha_word)
    assume(target.strip())  # ensure non-empty after stripping

    return token, target, pattern


@given(pair=token_and_target())
@settings(max_examples=200)
def test_glossary_match_case_capitalisation_preservation(pair):
    # Feature: translation-layout-quality, Property 6: Glossary capitalisation preservation
    """
    Property 6: Glossary capitalisation preservation

    For any matched token in the translated text, the capitalisation pattern
    of the matched token (all-caps, title-case, lowercase) shall be applied
    to the substituted target term.

    Specifically:
    - An all-caps token  → `_match_case` returns `target.upper()`
    - A title-case token → `_match_case` returns `target.title()`
    - A lowercase token  → `_match_case` returns `target.lower()`

    **Validates: Requirements 3.3**
    """
    token, target, pattern = pair

    result = Glossary_Store._match_case(token, target)

    if pattern == "upper":
        # Token is all-caps → result must be all-caps target.
        assert result == target.upper(), (
            f"Expected all-caps result for token={token!r}, target={target!r}.\n"
            f"  Got: {result!r}, expected: {target.upper()!r}"
        )
        # Verify the token really is all-caps (sanity check on the generator).
        assert token.isupper(), f"Token {token!r} should be all-caps."

    elif pattern == "title":
        # Token is title-case → result must be title-case target.
        assert result == target.title(), (
            f"Expected title-case result for token={token!r}, target={target!r}.\n"
            f"  Got: {result!r}, expected: {target.title()!r}"
        )
        # Verify the token really is title-case.
        assert token.istitle(), f"Token {token!r} should be title-case."

    else:  # pattern == "lower"
        # Token is lowercase → result must be lowercase target.
        assert result == target.lower(), (
            f"Expected lowercase result for token={token!r}, target={target!r}.\n"
            f"  Got: {result!r}, expected: {target.lower()!r}"
        )
        # Verify the token really is lowercase.
        assert token.islower(), f"Token {token!r} should be lowercase."


# ===========================================================================
# Property 3: Chunk minimum size
# Feature: translation-layout-quality, Property 3: Chunk minimum size
# ===========================================================================

import random as _random

# A strategy that generates a single "word" token (no whitespace, non-empty).
_chunk_word = st.text(
    alphabet=st.characters(
        whitelist_categories=("Lu", "Ll", "Nd"),
        whitelist_characters="-'",
    ),
    min_size=1,
    max_size=12,
)


def _build_text_with_boundaries(words: list, boundary_positions: list) -> str:
    """
    Attach a period to the words at the given 0-based indices so the
    resulting text has predictable sentence boundaries.
    """
    result = list(words)
    for pos in boundary_positions:
        if 0 <= pos < len(result):
            result[pos] = result[pos].rstrip(".!?") + "."
    return " ".join(result)


@given(
    # Generate a list of at least 5 word tokens.
    words=st.lists(_chunk_word, min_size=5, max_size=300),
    # Fraction of words to mark as sentence boundaries.
    boundary_fraction=st.floats(min_value=0.05, max_value=0.5),
    # Seed for deterministic boundary placement.
    seed=st.integers(min_value=0, max_value=10_000),
)
@settings(max_examples=200)
def test_chunk_minimum_size(words, boundary_fraction, seed):
    # Feature: translation-layout-quality, Property 3: Chunk minimum size
    """
    Property 3: Chunk minimum size

    For any block of text with at least 5 whitespace-delimited tokens, every
    chunk produced by Chunk_Splitter.split() shall contain at least 5
    whitespace-delimited tokens.

    **Validates: Requirements 1.4**
    """
    # Ensure we have at least 5 tokens (hypothesis min_size=5 guarantees this,
    # but be explicit for clarity).
    assume(len(words) >= 5)

    # Place sentence boundaries at roughly `boundary_fraction` of the words
    # so the splitter has real split points to work with.
    n_boundaries = max(1, int(len(words) * boundary_fraction))
    rng = _random.Random(seed)
    boundary_positions = rng.sample(range(len(words)), min(n_boundaries, len(words)))

    text = _build_text_with_boundaries(words, boundary_positions)

    splitter = Chunk_Splitter()
    chunks = splitter.split(text, max_tokens=80, hard_cap=150, min_tokens=5)

    # Every chunk must have at least 5 whitespace-delimited tokens.
    for i, chunk in enumerate(chunks):
        token_count = len(chunk.split())
        assert token_count >= 5, (
            f"Chunk {i} has only {token_count} token(s) (min=5).\n"
            f"  chunk: {chunk!r}\n"
            f"  full text ({len(words)} tokens): {text[:120]!r}..."
        )


# ===========================================================================
# Property 5: Glossary whole-word substitution
# Feature: translation-layout-quality, Property 5: Glossary whole-word substitution
# ===========================================================================

# Strategy: generate a single alphabetic word (no digits, no punctuation,
# no whitespace) to use as a glossary source or target term.
_glossary_word = st.text(
    alphabet=st.characters(whitelist_categories=("Lu", "Ll")),
    min_size=2,
    max_size=12,
).filter(lambda w: w.isalpha() and w.strip() == w)

# Strategy: generate a list of unique (source, target) pairs where source
# and target are distinct alphabetic words and no two sources are the same
# (case-insensitive).  We keep the list small (1–6 pairs) to stay fast.
@st.composite
def glossary_pairs_and_text(draw):
    """
    Draw a list of (source, target) glossary pairs and a text string that
    contains at least one whole-word occurrence of each source term, plus
    some "embedded" occurrences where the source term appears as a substring
    of a longer word (to verify partial-word non-replacement).

    Returns (pairs, text) where:
    - pairs: list of (source, target) with unique case-insensitive sources
    - text: a string built from the source terms (whole-word and embedded)
      plus filler words

    Constraints to avoid ambiguous test cases:
    - No source term is a substring of another source term (case-insensitive),
      which would cause the embedded form of one term to contain another term
      as a whole word.
    - No target term equals any source term (case-insensitive), which would
      cause circular substitutions where a replacement introduces a new match.
    - No source term is a substring of any target term and vice versa in a
      way that would cause cascading replacements.
    """
    n_pairs = draw(st.integers(min_value=1, max_value=4))

    # Draw n_pairs unique source terms (case-insensitive uniqueness).
    # Also ensure no source is a substring of another source (to avoid
    # embedded form collisions between different terms).
    sources = []
    seen_lower = set()
    attempts = 0
    while len(sources) < n_pairs and attempts < 300:
        attempts += 1
        word = draw(_glossary_word)
        wl = word.lower()
        # Skip if already seen (case-insensitive duplicate).
        if wl in seen_lower:
            continue
        # Skip if this word is a substring of any existing source or vice versa.
        if any(wl in s or s in wl for s in seen_lower):
            continue
        seen_lower.add(wl)
        sources.append(word)

    # If we couldn't get enough unique sources, use what we have.
    n_pairs = len(sources)
    assume(n_pairs >= 1)

    sources_lower = {s.lower() for s in sources}

    # Draw a distinct target for each source.
    # Constraints:
    # - Target must differ from its own source (case-insensitive).
    # - Target must not equal any source term (prevents circular substitution).
    # - Target must not be a substring of any source (prevents embedded match
    #   in the embedded form "pre<source>fix" after substitution).
    targets = []
    for src in sources:
        for _ in range(100):
            tgt = draw(_glossary_word)
            tl = tgt.lower()
            # Must differ from source.
            if tl == src.lower():
                continue
            # Must not equal any source term (no circular substitution).
            if tl in sources_lower:
                continue
            # Must not be a substring of any source (no partial match after
            # substitution creates a new whole-word match).
            if any(tl in sl or sl in tl for sl in sources_lower):
                continue
            targets.append(tgt)
            break
        else:
            # Could not find a valid target; skip this example.
            assume(False)

    pairs = list(zip(sources, targets))

    # Build a text string that contains:
    # 1. Each source term as a standalone whole word (surrounded by spaces).
    # 2. Each source term embedded inside a longer word (prefix + source + suffix)
    #    so we can verify partial-word occurrences are NOT replaced.
    filler_words = ["the", "quick", "brown", "fox", "jumps", "over", "lazy", "dog"]
    tokens = []
    for src in sources:
        # Whole-word occurrence.
        tokens.append(src)
        # Embedded occurrence: prepend "pre" and append "fix" to make a longer
        # word that contains the source as a substring but not as a whole word.
        embedded = "pre" + src + "fix"
        tokens.append(embedded)
        # Add a filler word for spacing variety.
        tokens.append(filler_words[len(tokens) % len(filler_words)])

    # Shuffle tokens so the order is random.
    seed = draw(st.integers(min_value=0, max_value=10_000))
    rng = _random.Random(seed)
    rng.shuffle(tokens)

    text = " ".join(tokens)
    return pairs, text


@given(data=glossary_pairs_and_text())
@settings(max_examples=100)
def test_glossary_whole_word_substitution(data):
    # Feature: translation-layout-quality, Property 5: Glossary whole-word substitution
    """
    Property 5: Glossary whole-word substitution

    After calling ``Glossary_Store.apply(text)``:

    1. No whole-word occurrence of any source term shall remain in the output
       (all whole-word matches must have been replaced).
    2. Partial-word occurrences of a source term (where the term appears as a
       substring of a longer word) shall NOT be replaced — the surrounding
       characters must be preserved unchanged.

    **Validates: Requirements 3.2**
    """
    pairs, text = data

    store = Glossary_Store(pairs)
    result = store.apply(text)

    for source, target in pairs:
        # ── Assertion 1: no whole-word occurrence of source remains ──────
        # Use the same \b word-boundary regex that Glossary_Store uses.
        whole_word_pattern = re.compile(
            r"\b" + re.escape(source) + r"\b", re.IGNORECASE
        )
        remaining_matches = whole_word_pattern.findall(result)
        assert not remaining_matches, (
            f"Whole-word occurrence of source term {source!r} still present "
            f"in result after apply().\n"
            f"  Remaining matches: {remaining_matches}\n"
            f"  Source text: {text!r}\n"
            f"  Result text: {result!r}"
        )

        # ── Assertion 2: partial-word occurrences were NOT replaced ───────
        # The text contains tokens of the form "pre<source>fix".  After
        # apply(), those tokens must still contain the source term as a
        # substring (i.e., the embedded occurrence was not touched).
        embedded_form = "pre" + source + "fix"
        embedded_pattern_in_text = re.compile(
            r"\bpre" + re.escape(source) + r"fix\b", re.IGNORECASE
        )
        # Only check if the embedded form was actually present in the input.
        if embedded_pattern_in_text.search(text):
            # The embedded token must still appear in the result (unchanged).
            assert embedded_pattern_in_text.search(result), (
                f"Partial-word occurrence 'pre{source}fix' was incorrectly "
                f"modified or removed by apply().\n"
                f"  Source text: {text!r}\n"
                f"  Result text: {result!r}"
            )
            # Additionally, the embedded token must NOT have been replaced
            # with a version containing the target term in the middle.
            embedded_replaced = re.compile(
                r"\bpre" + re.escape(target) + r"fix\b", re.IGNORECASE
            )
            assert not embedded_replaced.search(result), (
                f"Partial-word occurrence 'pre{source}fix' was incorrectly "
                f"replaced with 'pre{target}fix' by apply().\n"
                f"  Source text: {text!r}\n"
                f"  Result text: {result!r}"
            )


# ===========================================================================
# Property 7: Overflow resolution does not move other blocks
# Feature: translation-layout-quality, Property 7: Overflow resolution does not move other blocks
# ===========================================================================

import copy
from unittest.mock import MagicMock

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from document_translator_v3 import _resolve_overflow


def _make_mock_page_for_property(insert_side_effects):
    """Build a minimal fitz.Page mock for property tests."""
    page = MagicMock()
    page.rect = MagicMock()
    page.rect.height = 792.0
    page.rect.width = 612.0
    page.insert_textbox = MagicMock(side_effect=insert_side_effects)
    page.draw_rect = MagicMock()
    page.show_pdf_page = MagicMock()
    return page


# Strategy: generate a list of non-overlapping block bounding boxes.
# Each bbox is [x0, y0, x1, y1] with x0 < x1 and y0 < y1.
@st.composite
def page_layout(draw):
    """
    Generate a page layout with:
    - A list of 1–5 "other" blocks (non-overflowing, read-only).
    - One "overflowing" block with a fixed position.

    The overflowing block is always at y=50–100.
    Other blocks are placed below it (y > 100) to avoid overlap.
    """
    n_other = draw(st.integers(min_value=0, max_value=5))

    other_bboxes = []
    y_cursor = 110.0
    for _ in range(n_other):
        height = draw(st.floats(min_value=10.0, max_value=60.0))
        bbox = [50.0, y_cursor, 400.0, y_cursor + height]
        other_bboxes.append(bbox)
        y_cursor += height + draw(st.floats(min_value=5.0, max_value=30.0))

    # Overflowing block position (fixed for simplicity).
    x0, y0, x1, y1 = 50.0, 50.0, 400.0, 100.0

    # Initial overflow amount (negative remaining).
    initial_remaining = draw(st.floats(min_value=-100.0, max_value=-1.0))

    return {
        "x0": x0, "y0": y0, "x1": x1, "y1": y1,
        "other_bboxes": other_bboxes,
        "initial_remaining": initial_remaining,
    }


@given(layout=page_layout())
@settings(max_examples=200)
def test_overflow_resolution_does_not_move_other_blocks(layout):
    """
    **Validates: Requirements 6.5**

    Property 7: Overflow resolution does not move other blocks

    For any PDF page layout with multiple blocks where one block overflows,
    the position and dimensions of all non-overflowing blocks shall remain
    unchanged after overflow resolution.

    Specifically, the ``other_bboxes`` list passed to ``_resolve_overflow()``
    must be identical (element-by-element) after the call returns.
    """
    # Feature: translation-layout-quality, Property 7: Overflow resolution does not move other blocks

    other_bboxes = layout["other_bboxes"]
    # Deep-copy the original bboxes so we can compare after the call.
    original_bboxes = copy.deepcopy(other_bboxes)

    # Provide enough insert_textbox return values to cover all possible
    # retry attempts (font 12→6 = up to 7 retries, plus continuation box).
    # Use a mix of negative (overflow) and positive (fits) values so both
    # code paths are exercised.
    insert_returns = [-5.0] * 10 + [5.0] * 5  # eventually fits or cont box

    page = _make_mock_page_for_property(insert_returns)
    original_doc = MagicMock()

    _resolve_overflow(
        page=page,
        block_text="Some overflowing translated text.",
        x0=layout["x0"],
        y0=layout["y0"],
        x1=layout["x1"],
        y1=layout["y1"],
        font_size=12.0,
        resolved_font="helv",
        other_bboxes=other_bboxes,
        page_height=792.0,
        original_doc=original_doc,
        page_num=0,
        bg_color=(1.0, 1.0, 1.0),
        initial_remaining=layout["initial_remaining"],
    )

    # The other_bboxes list must be identical to the original.
    assert other_bboxes == original_bboxes, (
        f"other_bboxes were modified during overflow resolution.\n"
        f"  Before: {original_bboxes}\n"
        f"  After:  {other_bboxes}"
    )


# ===========================================================================
# Property 9: DOCX paragraph formatting preservation
# Feature: translation-layout-quality, Property 9: DOCX paragraph formatting preservation
# ===========================================================================

import tempfile
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from document_translator_v3 import write_docx

# Alignment values: LEFT=0, CENTER=1, RIGHT=2, JUSTIFY=3
_alignment_strategy = st.sampled_from([
    WD_ALIGN_PARAGRAPH.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY,
])

# Spacing values: Pt(6) to Pt(24) in integer steps
_spacing_strategy = st.integers(min_value=6, max_value=24).map(Pt)


@st.composite
def paragraph_formatting_data(draw):
    """
    Draw a list of 1–5 paragraph blocks, each with:
    - A random alignment value (WD_ALIGN_PARAGRAPH enum member)
    - A random space_before value (Pt object, 6–24pt)
    - A random space_after value (Pt object, 6–24pt)

    Returns a list of block dicts suitable for write_docx().
    """
    n = draw(st.integers(min_value=1, max_value=5))
    blocks = []
    for i in range(n):
        alignment = draw(_alignment_strategy)
        space_before = draw(_spacing_strategy)
        space_after = draw(_spacing_strategy)
        blocks.append({
            "type": "paragraph",
            "text": f"Sample paragraph text number {i + 1} for formatting test.",
            "style": {
                "style_name": "Normal",
                "alignment": alignment,
                "space_before": space_before,
                "space_after": space_after,
                "bold": False,
                "font_size": None,
                "font_color": None,
            },
        })
    return blocks


@given(blocks=paragraph_formatting_data())
@settings(max_examples=100)
def test_docx_paragraph_formatting_preservation(blocks):
    # Feature: translation-layout-quality, Property 9: DOCX paragraph formatting preservation
    """
    Property 9: DOCX paragraph formatting preservation

    For any list of paragraph blocks with explicitly set alignment and spacing
    values, the output DOCX produced by ``write_docx()`` must preserve those
    values exactly.

    Specifically:
    - ``para.alignment`` in the output must equal the source alignment value.
    - ``para.paragraph_format.space_before`` in the output must equal the
      source ``space_before`` value (in EMUs, as stored by python-docx).
    - ``para.paragraph_format.space_after`` in the output must equal the
      source ``space_after`` value (in EMUs, as stored by python-docx).

    **Validates: Requirements 7.4, 7.5**
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        write_docx(blocks, output_path)

        # Read back the output document and verify formatting.
        from docx import Document as _Document
        out_doc = _Document(output_path)
        out_paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]

        # The number of output paragraphs must match the number of input blocks.
        assert len(out_paragraphs) == len(blocks), (
            f"Expected {len(blocks)} paragraphs in output, "
            f"got {len(out_paragraphs)}."
        )

        for i, (block, out_para) in enumerate(zip(blocks, out_paragraphs)):
            src_style = block["style"]
            src_alignment = src_style["alignment"]
            src_space_before = src_style["space_before"]
            src_space_after = src_style["space_after"]

            # ── Alignment check (Requirement 7.4) ────────────────────────
            assert out_para.alignment == src_alignment, (
                f"Paragraph {i}: alignment mismatch.\n"
                f"  Expected: {src_alignment!r}\n"
                f"  Got:      {out_para.alignment!r}"
            )

            # ── space_before check (Requirement 7.5) ─────────────────────
            assert out_para.paragraph_format.space_before == src_space_before, (
                f"Paragraph {i}: space_before mismatch.\n"
                f"  Expected: {src_space_before!r}\n"
                f"  Got:      {out_para.paragraph_format.space_before!r}"
            )

            # ── space_after check (Requirement 7.5) ──────────────────────
            assert out_para.paragraph_format.space_after == src_space_after, (
                f"Paragraph {i}: space_after mismatch.\n"
                f"  Expected: {src_space_after!r}\n"
                f"  Got:      {out_para.paragraph_format.space_after!r}"
            )
    finally:
        # Clean up the temporary file.
        try:
            os.remove(output_path)
        except OSError:
            pass


# ===========================================================================
# Property 8: DOCX paragraph style round-trip
# Feature: translation-layout-quality, Property 8: DOCX paragraph style round-trip
# ===========================================================================

# Standard DOCX styles that are available in a default python-docx Document.
_STANDARD_DOCX_STYLES = [
    "Normal",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "List Bullet",
    "List Number",
    "Body Text",
    "Caption",
]

# Strategy: generate a non-empty list of style names drawn from the standard
# DOCX styles set.  Each entry in the list corresponds to one paragraph block.
_style_name_list = st.lists(
    st.sampled_from(_STANDARD_DOCX_STYLES),
    min_size=1,
    max_size=10,
)


@given(style_names=_style_name_list)
@settings(max_examples=100)
def test_docx_paragraph_style_round_trip(style_names):
    # Feature: translation-layout-quality, Property 8: DOCX paragraph style round-trip
    """
    Property 8: DOCX paragraph style round-trip

    For any list of paragraph style names drawn from the set of standard DOCX
    styles, the output DOCX produced by ``write_docx()`` must apply the same
    paragraph style name to each output paragraph whose style exists in the
    output document's style table.

    Specifically, for every paragraph block whose ``style_name`` is present in
    the output document's available styles, the output paragraph's
    ``para.style.name`` must equal the source ``style_name``.

    **Validates: Requirements 7.2**
    """
    # Build paragraph blocks — one per style name in the generated list.
    blocks = []
    for i, style_name in enumerate(style_names):
        blocks.append({
            "type": "paragraph",
            "text": f"Sample paragraph {i + 1} with style {style_name!r}.",
            "style": {
                "style_name": style_name,
                "alignment": None,
                "space_before": None,
                "space_after": None,
                "bold": False,
                "font_size": None,
                "font_color": None,
            },
        })

    # Write to a temporary DOCX file and read it back.
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    output_path = tmp.name
    tmp.close()

    try:
        write_docx(blocks, output_path)

        # Read back the output document.
        from docx import Document as _Document
        out_doc = _Document(output_path)

        # Build the set of style names available in the output document.
        available_styles = {s.name for s in out_doc.styles}

        # Collect the output paragraphs (skip empty default paragraphs that
        # python-docx may add; match by index against the input blocks).
        out_paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]

        # There must be at least as many output paragraphs as input blocks.
        assert len(out_paragraphs) >= len(blocks), (
            f"Expected at least {len(blocks)} non-empty paragraphs in output, "
            f"got {len(out_paragraphs)}.\n"
            f"  style_names: {style_names!r}"
        )

        # For each input block, verify the style round-trip for styles that
        # exist in the output document.
        for i, (block, out_para) in enumerate(zip(blocks, out_paragraphs)):
            src_style_name = block["style"]["style_name"]

            # Only assert round-trip when the style exists in the output doc.
            if src_style_name in available_styles:
                assert out_para.style.name == src_style_name, (
                    f"Paragraph {i}: style name mismatch.\n"
                    f"  Expected: {src_style_name!r}\n"
                    f"  Got:      {out_para.style.name!r}\n"
                    f"  style_names: {style_names!r}"
                )

    finally:
        # Clean up the temporary file.
        try:
            os.remove(output_path)
        except OSError:
            pass


# ===========================================================================
# Property 10: DOCX run font color preservation
# Feature: translation-layout-quality, Property 10: DOCX run font color preservation
# ===========================================================================

from docx.shared import RGBColor


@st.composite
def font_color_blocks(draw):
    """
    Draw a list of 1–5 paragraph blocks, each with a random RGBColor value
    (r, g, b each in 0–255) set as the run font color.

    Returns a list of block dicts suitable for write_docx().
    """
    n = draw(st.integers(min_value=1, max_value=5))
    blocks = []
    for i in range(n):
        r = draw(st.integers(min_value=0, max_value=255))
        g = draw(st.integers(min_value=0, max_value=255))
        b = draw(st.integers(min_value=0, max_value=255))
        color = RGBColor(r, g, b)
        blocks.append({
            "type": "paragraph",
            "text": f"Sample paragraph {i + 1} with font color ({r}, {g}, {b}).",
            "style": {
                "style_name": "Normal",
                "alignment": None,
                "space_before": None,
                "space_after": None,
                "bold": False,
                "font_size": None,
                "font_color": color,
            },
        })
    return blocks


@given(blocks=font_color_blocks())
@settings(max_examples=100)
def test_docx_run_font_color_preservation(blocks):
    # Feature: translation-layout-quality, Property 10: DOCX run font color preservation
    """
    Property 10: DOCX run font color preservation

    For any list of paragraph blocks with an explicitly set run font color
    (RGBColor), the output DOCX produced by ``write_docx()`` must preserve
    that color exactly on the corresponding output run.

    Specifically, for every paragraph block whose ``font_color`` is set,
    the first run of the output paragraph's ``run.font.color.rgb`` must equal
    the source ``font_color`` value.

    **Validates: Requirements 7.6**
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        write_docx(blocks, output_path)

        # Read back the output document and verify font colors.
        from docx import Document as _Document
        out_doc = _Document(output_path)
        out_paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]

        # The number of output paragraphs must match the number of input blocks.
        assert len(out_paragraphs) == len(blocks), (
            f"Expected {len(blocks)} paragraphs in output, "
            f"got {len(out_paragraphs)}."
        )

        for i, (block, out_para) in enumerate(zip(blocks, out_paragraphs)):
            src_color = block["style"]["font_color"]

            # The output paragraph must have at least one run.
            assert out_para.runs, (
                f"Paragraph {i}: no runs found in output paragraph."
            )

            out_run = out_para.runs[0]
            out_color = out_run.font.color.rgb

            assert out_color == src_color, (
                f"Paragraph {i}: font color mismatch.\n"
                f"  Expected: {src_color!r}\n"
                f"  Got:      {out_color!r}"
            )
    finally:
        # Clean up the temporary file.
        try:
            os.remove(output_path)
        except OSError:
            pass


# ===========================================================================
# Property 11: Table cell translation completeness
# Feature: translation-layout-quality, Property 11: Table cell translation completeness
# ===========================================================================

# Strategy: generate a single table cell block with either empty or non-empty text.
@st.composite
def table_with_mixed_cells(draw):
    """
    Draw a random table (1–5 rows × 1–5 cols) where each cell is either
    empty (text="") or non-empty (a short alphabetic string).

    Returns a list of table_cell block dicts suitable for write_docx().
    """
    num_rows = draw(st.integers(min_value=1, max_value=5))
    num_cols = draw(st.integers(min_value=1, max_value=5))

    # Ensure at least one non-empty cell so the property is meaningful.
    cells = []
    has_nonempty = False

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            # Draw whether this cell is empty or non-empty.
            is_empty = draw(st.booleans())

            if is_empty:
                cell_text = ""
            else:
                # Generate a non-empty cell text (at least 1 word).
                words = draw(st.lists(
                    st.text(
                        alphabet=st.characters(whitelist_categories=("Ll", "Lu")),
                        min_size=1,
                        max_size=10,
                    ).filter(lambda w: w.strip()),
                    min_size=1,
                    max_size=5,
                ))
                cell_text = " ".join(words)
                has_nonempty = True

            cells.append({
                "type": "table_cell",
                "text": cell_text,
                "table_index": 0,
                "row": row_idx,
                "col": col_idx,
                "row_span": 1,
                "col_span": 1,
                "style": {
                    "bold": None,
                    "font_size": None,
                },
            })

    # If all cells happened to be empty, force the first cell to be non-empty.
    if not has_nonempty:
        cells[0]["text"] = "NonEmptyCell"

    return cells


@given(cells=table_with_mixed_cells())
@settings(max_examples=100)
def test_table_cell_translation_completeness(cells):
    # Feature: translation-layout-quality, Property 11: Table cell translation completeness
    """
    Property 11: Table cell translation completeness

    For any table with a mix of empty and non-empty cells, ``write_docx()``
    must:
    1. Write non-empty cell text to the output (the cell text is present in
       the output document).
    2. Leave empty cells empty in the output (no text is written to cells
       whose source text is empty after stripping).

    Note: ``write_docx()`` writes pre-translated blocks directly, so this
    property verifies that the writer faithfully transfers non-empty cell
    content and does not populate empty cells.

    **Validates: Requirements 8.2**
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        write_docx(cells, output_path)

        # Read back the output document.
        from docx import Document as _Document
        out_doc = _Document(output_path)

        # The output document must contain exactly one table.
        assert len(out_doc.tables) == 1, (
            f"Expected 1 table in output, got {len(out_doc.tables)}."
        )

        out_table = out_doc.tables[0]

        # Build a lookup of source cells: (row, col) → stripped text.
        src_map = {(c["row"], c["col"]): c["text"].strip() for c in cells}

        num_rows = max(c["row"] for c in cells) + 1
        num_cols = max(c["col"] for c in cells) + 1

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                src_text = src_map.get((row_idx, col_idx), "")
                out_cell = out_table.cell(row_idx, col_idx)
                out_text = out_cell.text.strip()

                if src_text:
                    # Non-empty source cell: output must contain the text.
                    assert out_text == src_text, (
                        f"Cell ({row_idx}, {col_idx}): non-empty source text "
                        f"not written to output.\n"
                        f"  Source text: {src_text!r}\n"
                        f"  Output text: {out_text!r}"
                    )
                else:
                    # Empty source cell: output must also be empty.
                    assert out_text == "", (
                        f"Cell ({row_idx}, {col_idx}): empty source cell has "
                        f"non-empty output text.\n"
                        f"  Output text: {out_text!r}"
                    )
    finally:
        try:
            os.remove(output_path)
        except OSError:
            pass


# ===========================================================================
# Property 12: Table structural preservation
# Feature: translation-layout-quality, Property 12: Table structural preservation
# ===========================================================================


@st.composite
def table_dimensions_data(draw):
    """
    Draw random table dimensions (rows 1–10, cols 1–10) and build a list of
    table_cell block dicts with those dimensions.

    Returns a dict with:
    - "num_rows": int
    - "num_cols": int
    - "cells": list of table_cell block dicts
    """
    num_rows = draw(st.integers(min_value=1, max_value=10))
    num_cols = draw(st.integers(min_value=1, max_value=10))

    cells = []
    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            cell_text = f"Cell {row_idx},{col_idx}"
            cells.append({
                "type": "table_cell",
                "text": cell_text,
                "table_index": 0,
                "row": row_idx,
                "col": col_idx,
                "row_span": 1,
                "col_span": 1,
                "style": {
                    "bold": None,
                    "font_size": None,
                },
            })

    return {"num_rows": num_rows, "num_cols": num_cols, "cells": cells}


@given(data=table_dimensions_data())
@settings(max_examples=50)
def test_table_structural_preservation(data):
    # Feature: translation-layout-quality, Property 12: Table structural preservation
    """
    Property 12: Table structural preservation

    For any table with random dimensions (rows 1–10, cols 1–10), the output
    DOCX produced by ``write_docx()`` must reconstruct a table with the same
    number of rows and columns as the source.

    Specifically:
    - ``len(out_table.rows)`` must equal the source ``num_rows``.
    - ``len(out_table.columns)`` must equal the source ``num_cols``.

    **Validates: Requirements 8.3**
    """
    num_rows = data["num_rows"]
    num_cols = data["num_cols"]
    cells = data["cells"]

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        write_docx(cells, output_path)

        # Read back the output document.
        from docx import Document as _Document
        out_doc = _Document(output_path)

        # The output document must contain exactly one table.
        assert len(out_doc.tables) == 1, (
            f"Expected 1 table in output, got {len(out_doc.tables)}.\n"
            f"  Source dimensions: {num_rows} rows × {num_cols} cols"
        )

        out_table = out_doc.tables[0]

        # Row count must match.
        out_num_rows = len(out_table.rows)
        assert out_num_rows == num_rows, (
            f"Row count mismatch.\n"
            f"  Expected: {num_rows}\n"
            f"  Got:      {out_num_rows}"
        )

        # Column count must match.
        out_num_cols = len(out_table.columns)
        assert out_num_cols == num_cols, (
            f"Column count mismatch.\n"
            f"  Expected: {num_cols}\n"
            f"  Got:      {out_num_cols}"
        )

    finally:
        try:
            os.remove(output_path)
        except OSError:
            pass


# ===========================================================================
# Property 13: Table run formatting preservation
# Feature: translation-layout-quality, Property 13: Table run formatting preservation
# ===========================================================================


@st.composite
def table_run_formatting_data(draw):
    """
    Draw a list of table_cell block dicts (1–5 cells in a single-row table)
    where each cell has:
    - A random bold value: True, False, or None (unset)
    - A random font_size value: None (unset) or a Pt value in 6–72pt

    Returns a list of table_cell block dicts suitable for write_docx().
    """
    num_cols = draw(st.integers(min_value=1, max_value=5))

    cells = []
    for col_idx in range(num_cols):
        # bold: True, False, or None
        bold = draw(st.one_of(st.just(True), st.just(False), st.none()))

        # font_size: None or Pt(6)–Pt(72)
        font_size = draw(
            st.one_of(
                st.none(),
                st.integers(min_value=6, max_value=72).map(Pt),
            )
        )

        cells.append({
            "type": "table_cell",
            "text": f"Cell content {col_idx}",
            "table_index": 0,
            "row": 0,
            "col": col_idx,
            "row_span": 1,
            "col_span": 1,
            "style": {
                "bold": bold,
                "font_size": font_size,
            },
        })

    return cells


@given(cells=table_run_formatting_data())
@settings(max_examples=100)
def test_table_run_formatting_preservation(cells):
    # Feature: translation-layout-quality, Property 13: Table run formatting preservation
    """
    Property 13: Table run formatting preservation

    For any table cell whose run has bold and/or font size explicitly set,
    the output DOCX produced by ``write_docx()`` must:
    1. Set bold on the output run when the source bold is True or False
       (explicitly set); leave bold unset (None) when source bold is None.
    2. Set font size on the output run when the source font_size is explicitly
       set (a Pt value); leave font size unset (None) when source is None.

    **Validates: Requirements 8.4**
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        write_docx(cells, output_path)

        # Read back the output document.
        from docx import Document as _Document
        out_doc = _Document(output_path)

        # The output document must contain exactly one table.
        assert len(out_doc.tables) == 1, (
            f"Expected 1 table in output, got {len(out_doc.tables)}."
        )

        out_table = out_doc.tables[0]

        for cell_block in cells:
            col_idx = cell_block["col"]
            src_bold = cell_block["style"]["bold"]
            src_font_size = cell_block["style"]["font_size"]

            out_cell = out_table.cell(0, col_idx)

            # Collect the first run from the output cell (across all paragraphs).
            out_run = None
            for para in out_cell.paragraphs:
                if para.runs:
                    out_run = para.runs[0]
                    break

            # If the source cell has non-empty text, there must be a run.
            if cell_block["text"].strip():
                assert out_run is not None, (
                    f"Cell (0, {col_idx}): no run found in output cell "
                    f"for non-empty source text {cell_block['text']!r}."
                )

            if out_run is None:
                # Empty cell — no run to check formatting on.
                continue

            # ── Bold check (Requirement 8.4) ─────────────────────────────
            out_bold = out_run.bold
            if src_bold is not None:
                # Explicitly set in source → must match in output.
                assert out_bold == src_bold, (
                    f"Cell (0, {col_idx}): bold mismatch.\n"
                    f"  Source bold: {src_bold!r}\n"
                    f"  Output bold: {out_bold!r}"
                )
            else:
                # Not set in source → must be unset (None) in output.
                assert out_bold is None, (
                    f"Cell (0, {col_idx}): bold should be unset (None) "
                    f"when source bold is None, but got {out_bold!r}."
                )

            # ── Font size check (Requirement 8.4) ────────────────────────
            out_font_size = out_run.font.size
            if src_font_size is not None:
                # Explicitly set in source → must match in output (in EMUs).
                assert out_font_size == src_font_size, (
                    f"Cell (0, {col_idx}): font_size mismatch.\n"
                    f"  Source font_size: {src_font_size!r}\n"
                    f"  Output font_size: {out_font_size!r}"
                )
            else:
                # Not set in source → must be unset (None) in output.
                assert out_font_size is None, (
                    f"Cell (0, {col_idx}): font_size should be unset (None) "
                    f"when source font_size is None, but got {out_font_size!r}."
                )

    finally:
        try:
            os.remove(output_path)
        except OSError:
            pass

# ===========================================================================
# Property 14: Column gap threshold
# Feature: translation-layout-quality, Property 14: Column gap threshold
# ===========================================================================

from document_translator_v3 import detect_columns


@st.composite
def two_group_layout(draw):
    """
    Generate a page layout with two groups of blocks (left and right) and a
    controlled horizontal gap between them.

    The left group occupies x-range [left_x0, left_x1] and the right group
    occupies x-range [right_x0, right_x1], where:
        gap = right_x0 - left_x1

    To avoid floating-point precision issues near the 10% threshold, the gap
    is chosen to be either clearly above (>= 15%) or clearly below (<= 5%)
    the 10% threshold, keeping a 5% margin on each side.

    Returns a dict with:
    - "page_width": float  (100–1000)
    - "blocks": list of block dicts
    - "gap": float  (the actual gap between the two groups)
    - "gap_fraction": float  (gap / page_width)
    - "gap_meets_threshold": bool  (gap >= 10% of page_width)
    """
    page_width = float(draw(st.integers(min_value=100, max_value=1000)))

    # Decide whether the gap should clearly meet the threshold or clearly not.
    # Use a 5% margin on each side of the 10% threshold to avoid floating-point
    # ambiguity: "above" means >= 15%, "below" means <= 5%.
    gap_meets_threshold = draw(st.booleans())

    if gap_meets_threshold:
        # gap in [15%, 30%] of page_width — clearly above the 10% threshold
        gap_pct = draw(st.integers(min_value=15, max_value=30))
    else:
        # gap in [0%, 5%] of page_width — clearly below the 10% threshold
        gap_pct = draw(st.integers(min_value=0, max_value=5))

    # Compute gap as an integer percentage of page_width to keep arithmetic
    # exact and avoid floating-point rounding near the threshold.
    gap = round(gap_pct * page_width / 100.0, 6)

    # Left group: x0 = 0, width = 20% of page_width (fixed, simple)
    left_x0 = 0.0
    left_x1 = round(0.20 * page_width, 6)

    # Right group starts at left_x1 + gap
    right_x0 = round(left_x1 + gap, 6)
    right_x1 = round(right_x0 + 0.20 * page_width, 6)

    # Ensure right group fits within page width (skip if not).
    assume(right_x1 <= page_width)

    # Ensure block widths are strictly less than 90% of page_width so
    # blocks are NOT classified as full-width (which would bypass column logic).
    # With 20% width, this is always satisfied, but be explicit.
    assume((left_x1 - left_x0) < 0.90 * page_width)
    assume((right_x1 - right_x0) < 0.90 * page_width)

    # Generate >= 2 blocks per group (to avoid single-column fallback).
    n_left = draw(st.integers(min_value=2, max_value=5))
    n_right = draw(st.integers(min_value=2, max_value=5))

    blocks = []
    y_cursor = 10.0

    for i in range(n_left):
        block_height = draw(st.floats(min_value=5.0, max_value=30.0))
        blocks.append({
            "position": [left_x0, y_cursor, left_x1, y_cursor + block_height],
            "text": f"Left block {i}",
            "page": 0,
        })
        y_cursor += block_height + 2.0

    y_cursor = 10.0  # reset y for right column (independent y positions)

    for i in range(n_right):
        block_height = draw(st.floats(min_value=5.0, max_value=30.0))
        blocks.append({
            "position": [right_x0, y_cursor, right_x1, y_cursor + block_height],
            "text": f"Right block {i}",
            "page": 0,
        })
        y_cursor += block_height + 2.0

    # Compute the actual gap from block positions (for reporting).
    actual_gap = right_x0 - left_x1

    return {
        "page_width": page_width,
        "blocks": blocks,
        "gap": actual_gap,
        "gap_fraction": actual_gap / page_width,
        "gap_meets_threshold": gap_meets_threshold,
    }


@given(layout=two_group_layout())
@settings(max_examples=200)
def test_column_gap_threshold(layout):
    """
    **Validates: Requirements 10.1**

    Property 14: Column gap threshold

    For any page layout with two groups of blocks separated by a controlled
    horizontal gap:

    - WHEN the gap is >= 10% of the page width, ``detect_columns()`` SHALL
      detect a column boundary and return blocks in two-column reading order
      (all left-column blocks before all right-column blocks).
    - WHEN the gap is < 10% of the page width, ``detect_columns()`` SHALL
      treat the page as single-column and return all blocks sorted by top
      y-coordinate.

    Each group has >= 2 blocks to avoid the single-column fallback triggered
    by insufficient blocks per candidate column.
    """
    # Feature: translation-layout-quality, Property 14: Column gap threshold

    page_width = layout["page_width"]
    blocks = layout["blocks"]
    gap_meets_threshold = layout["gap_meets_threshold"]

    result = detect_columns(blocks, page_width)

    # All input blocks must appear in the output (no blocks dropped).
    assert len(result) == len(blocks), (
        f"detect_columns() returned {len(result)} blocks but input had "
        f"{len(blocks)} blocks.\n"
        f"  page_width={page_width:.2f}, gap={layout['gap']:.2f} "
        f"({layout['gap_fraction']*100:.1f}% of page_width)"
    )

    # Identify left-group and right-group blocks by their text prefix.
    left_texts = {b["text"] for b in blocks if b["text"].startswith("Left")}
    right_texts = {b["text"] for b in blocks if b["text"].startswith("Right")}

    result_texts = [b["text"] for b in result]

    if gap_meets_threshold:
        # Column boundary detected: all left-column blocks must appear before
        # all right-column blocks in the output.
        left_indices = [i for i, t in enumerate(result_texts) if t in left_texts]
        right_indices = [i for i, t in enumerate(result_texts) if t in right_texts]

        assert left_indices and right_indices, (
            f"Expected both left and right blocks in result.\n"
            f"  result_texts: {result_texts}\n"
            f"  page_width={page_width:.2f}, gap={layout['gap']:.2f} "
            f"({layout['gap_fraction']*100:.1f}% of page_width)"
        )

        max_left_idx = max(left_indices)
        min_right_idx = min(right_indices)

        assert max_left_idx < min_right_idx, (
            f"Column boundary detected but left blocks do not all precede "
            f"right blocks in the output.\n"
            f"  left_indices={left_indices}, right_indices={right_indices}\n"
            f"  result_texts={result_texts}\n"
            f"  page_width={page_width:.2f}, gap={layout['gap']:.2f} "
            f"({layout['gap_fraction']*100:.1f}% of page_width)"
        )
    else:
        # No column boundary: blocks returned sorted by top y-coordinate.
        result_y_tops = [b["position"][1] for b in result]
        assert result_y_tops == sorted(result_y_tops), (
            f"Single-column fallback: blocks not sorted by top y-coordinate.\n"
            f"  y_tops={result_y_tops}\n"
            f"  page_width={page_width:.2f}, gap={layout['gap']:.2f} "
            f"({layout['gap_fraction']*100:.1f}% of page_width)"
        )


# ===========================================================================
# Property 15: Column reading order
# Feature: translation-layout-quality, Property 15: Column reading order
# ===========================================================================


@st.composite
def two_column_layout_shuffled(draw):
    """
    Generate a two-column page layout with shuffled y-coordinates.

    The layout has:
    - A clear 15%+ gap between the left and right column groups (so the
      column boundary is always detected).
    - Each column has >= 2 blocks with distinct, randomly ordered y-positions.
    - Blocks are shuffled before being passed to detect_columns() to verify
      that the function correctly sorts them by y within each column.

    Returns a dict with:
    - "page_width": float
    - "blocks": list of block dicts (shuffled order)
    - "left_block_texts": set of text labels for left-column blocks
    - "right_block_texts": set of text labels for right-column blocks
    - "left_y_tops": list of y0 values for left-column blocks (in generation order)
    - "right_y_tops": list of y0 values for right-column blocks (in generation order)
    """
    page_width = float(draw(st.integers(min_value=200, max_value=800)))

    # Left column: x in [0, 25% of page_width]
    left_x0 = 0.0
    left_x1 = round(0.25 * page_width, 4)

    # Gap: 15%–25% of page_width (clearly above the 10% threshold)
    gap_pct = draw(st.integers(min_value=15, max_value=25))
    gap = round(gap_pct * page_width / 100.0, 4)

    # Right column: starts at left_x1 + gap
    right_x0 = round(left_x1 + gap, 4)
    right_x1 = round(right_x0 + 0.25 * page_width, 4)

    # Ensure right column fits within page width.
    assume(right_x1 <= page_width)

    # Ensure neither column is classified as full-width (< 90% of page_width).
    assume((left_x1 - left_x0) < 0.90 * page_width)
    assume((right_x1 - right_x0) < 0.90 * page_width)

    # Generate >= 2 blocks per column with distinct y-positions.
    n_left = draw(st.integers(min_value=2, max_value=6))
    n_right = draw(st.integers(min_value=2, max_value=6))

    # Draw distinct y0 values for left column blocks (0–700 range).
    left_y_tops = draw(
        st.lists(
            st.floats(min_value=0.0, max_value=700.0, allow_nan=False, allow_infinity=False),
            min_size=n_left,
            max_size=n_left,
            unique=True,
        )
    )
    # Draw distinct y0 values for right column blocks (0–700 range).
    right_y_tops = draw(
        st.lists(
            st.floats(min_value=0.0, max_value=700.0, allow_nan=False, allow_infinity=False),
            min_size=n_right,
            max_size=n_right,
            unique=True,
        )
    )

    block_height = 20.0  # fixed height for simplicity

    left_blocks = []
    for i, y0 in enumerate(left_y_tops):
        left_blocks.append({
            "position": [left_x0, y0, left_x1, y0 + block_height],
            "text": f"Left_{i}",
            "page": 0,
        })

    right_blocks = []
    for i, y0 in enumerate(right_y_tops):
        right_blocks.append({
            "position": [right_x0, y0, right_x1, y0 + block_height],
            "text": f"Right_{i}",
            "page": 0,
        })

    # Shuffle all blocks together to simulate unsorted input.
    all_blocks = left_blocks + right_blocks
    seed = draw(st.integers(min_value=0, max_value=100_000))
    import random as _rand
    rng = _rand.Random(seed)
    rng.shuffle(all_blocks)

    return {
        "page_width": page_width,
        "blocks": all_blocks,
        "left_block_texts": {b["text"] for b in left_blocks},
        "right_block_texts": {b["text"] for b in right_blocks},
        "left_y_tops": left_y_tops,
        "right_y_tops": right_y_tops,
    }


@given(layout=two_column_layout_shuffled())
@settings(max_examples=200)
def test_column_reading_order(layout):
    # Feature: translation-layout-quality, Property 15: Column reading order
    """
    Property 15: Column reading order

    For any two-column page layout with shuffled y-coordinates and a clear
    15%+ gap between columns:

    1. Within each column group, blocks must be sorted by their top
       y-coordinate (ascending) in the output of ``detect_columns()``.
    2. All left-column blocks must appear before all right-column blocks in
       the output (left-to-right column ordering).

    **Validates: Requirements 10.3**
    """
    page_width = layout["page_width"]
    blocks = layout["blocks"]
    left_texts = layout["left_block_texts"]
    right_texts = layout["right_block_texts"]

    result = detect_columns(blocks, page_width)

    # All input blocks must appear in the output (no blocks dropped).
    assert len(result) == len(blocks), (
        f"detect_columns() returned {len(result)} blocks but input had "
        f"{len(blocks)} blocks.\n"
        f"  page_width={page_width}"
    )

    # Partition the result into left and right column sequences (preserving
    # the order they appear in the output).
    result_left = [b for b in result if b["text"] in left_texts]
    result_right = [b for b in result if b["text"] in right_texts]

    # ── Assertion 1: within each column, blocks are sorted by top y ──────────
    left_y_in_result = [b["position"][1] for b in result_left]
    assert left_y_in_result == sorted(left_y_in_result), (
        f"Left column blocks are not sorted by top y-coordinate.\n"
        f"  y_tops in result: {left_y_in_result}\n"
        f"  expected sorted:  {sorted(left_y_in_result)}"
    )

    right_y_in_result = [b["position"][1] for b in result_right]
    assert right_y_in_result == sorted(right_y_in_result), (
        f"Right column blocks are not sorted by top y-coordinate.\n"
        f"  y_tops in result: {right_y_in_result}\n"
        f"  expected sorted:  {sorted(right_y_in_result)}"
    )

    # ── Assertion 2: left column appears before right column ─────────────────
    result_texts = [b["text"] for b in result]
    left_indices = [i for i, b in enumerate(result) if b["text"] in left_texts]
    right_indices = [i for i, b in enumerate(result) if b["text"] in right_texts]

    assert left_indices and right_indices, (
        f"Expected both left and right blocks in result.\n"
        f"  result_texts: {result_texts}"
    )

    max_left_idx = max(left_indices)
    min_right_idx = min(right_indices)

    assert max_left_idx < min_right_idx, (
        f"Left-column blocks do not all precede right-column blocks.\n"
        f"  left_indices={left_indices}, right_indices={right_indices}\n"
        f"  result_texts={result_texts}\n"
        f"  page_width={page_width}"
    )


# ===========================================================================
# Property 16: Full-width block ordering
# Feature: translation-layout-quality, Property 16: Full-width block ordering
# ===========================================================================


@st.composite
def full_width_and_columnar_layout(draw):
    """
    Generate a page with a mix of full-width blocks and columnar blocks at
    various y-positions.

    Full-width blocks have width >= 90% of page_width.
    Columnar blocks have width < 90% of page_width (and are narrow enough
    to form at least two columns with a clear gap, so the single-column
    fallback is not triggered).

    Returns a dict with:
    - "page_width": float
    - "blocks": list of block dicts (shuffled order)
    - "full_width_texts": set of text labels for full-width blocks
    - "columnar_texts": set of text labels for columnar blocks
    """
    page_width = float(draw(st.integers(min_value=200, max_value=800)))

    # ── Full-width blocks ────────────────────────────────────────────────────
    # Width is in [90%, 100%] of page_width (full-width threshold is 90%).
    n_fw = draw(st.integers(min_value=1, max_value=4))

    fw_y_tops = draw(
        st.lists(
            st.floats(min_value=0.0, max_value=700.0, allow_nan=False, allow_infinity=False),
            min_size=n_fw,
            max_size=n_fw,
            unique=True,
        )
    )

    fw_width_pct = draw(
        st.lists(
            # Use 91%–100% (not exactly 90%) to avoid floating-point precision
            # issues where 0.9 * page_width may be slightly larger than
            # round(90 * page_width / 100, 4) due to IEEE 754 rounding.
            # A 1% margin ensures the block width is always strictly above the
            # 90% threshold used by detect_columns().
            st.integers(min_value=91, max_value=100),
            min_size=n_fw,
            max_size=n_fw,
        )
    )

    block_height = 15.0  # fixed height for simplicity

    full_width_blocks = []
    for i, (y0, w_pct) in enumerate(zip(fw_y_tops, fw_width_pct)):
        fw_width = round(w_pct * page_width / 100.0, 4)
        full_width_blocks.append({
            "position": [0.0, y0, fw_width, y0 + block_height],
            "text": f"FW_{i}",
            "page": 0,
        })

    # ── Columnar blocks ──────────────────────────────────────────────────────
    # Two narrow columns with a clear 15%+ gap so detect_columns() detects
    # two columns (avoiding the single-column fallback).
    # Left column: x in [0, 20% of page_width]
    left_x0 = 0.0
    left_x1 = round(0.20 * page_width, 4)

    # Gap: 15% of page_width (clearly above the 10% threshold)
    gap = round(0.15 * page_width, 4)

    # Right column: starts at left_x1 + gap
    right_x0 = round(left_x1 + gap, 4)
    right_x1 = round(right_x0 + 0.20 * page_width, 4)

    # Ensure right column fits within page width.
    assume(right_x1 <= page_width)

    # Ensure neither column is classified as full-width (< 90% of page_width).
    assume((left_x1 - left_x0) < 0.90 * page_width)
    assume((right_x1 - right_x0) < 0.90 * page_width)

    # Generate >= 2 blocks per column (to avoid single-column fallback).
    n_left = draw(st.integers(min_value=2, max_value=4))
    n_right = draw(st.integers(min_value=2, max_value=4))

    col_y_tops = draw(
        st.lists(
            st.floats(min_value=0.0, max_value=700.0, allow_nan=False, allow_infinity=False),
            min_size=n_left + n_right,
            max_size=n_left + n_right,
            unique=True,
        )
    )
    left_y_tops = col_y_tops[:n_left]
    right_y_tops = col_y_tops[n_left:]

    columnar_blocks = []
    for i, y0 in enumerate(left_y_tops):
        columnar_blocks.append({
            "position": [left_x0, y0, left_x1, y0 + block_height],
            "text": f"Col_L_{i}",
            "page": 0,
        })
    for i, y0 in enumerate(right_y_tops):
        columnar_blocks.append({
            "position": [right_x0, y0, right_x1, y0 + block_height],
            "text": f"Col_R_{i}",
            "page": 0,
        })

    # Shuffle all blocks together to simulate unsorted input.
    all_blocks = full_width_blocks + columnar_blocks
    seed = draw(st.integers(min_value=0, max_value=100_000))
    import random as _rand
    rng = _rand.Random(seed)
    rng.shuffle(all_blocks)

    return {
        "page_width": page_width,
        "blocks": all_blocks,
        "full_width_texts": {b["text"] for b in full_width_blocks},
        "columnar_texts": {b["text"] for b in columnar_blocks},
        # Map full-width block text → its top y-coordinate (for assertion).
        "fw_y_by_text": {b["text"]: b["position"][1] for b in full_width_blocks},
    }


@given(layout=full_width_and_columnar_layout())
@settings(max_examples=200)
def test_full_width_block_ordering(layout):
    # Feature: translation-layout-quality, Property 16: Full-width block ordering
    """
    **Validates: Requirements 10.4**

    Property 16: Full-width block ordering

    For any page containing full-width blocks (width >= 90% of page_width)
    and columnar blocks (width < 90% of page_width) at various y-positions,
    ``detect_columns()`` must place each full-width block before every
    columnar block whose top y-coordinate is strictly greater than the
    full-width block's top y-coordinate.

    Formally, for every full-width block F and every columnar block C in the
    output:
        IF C["position"][1] > F["position"][1]
        THEN index_of(F, result) < index_of(C, result)

    This ensures full-width blocks act as "section dividers" that appear
    before any columnar content that starts below them.
    """
    page_width = layout["page_width"]
    blocks = layout["blocks"]
    full_width_texts = layout["full_width_texts"]
    columnar_texts = layout["columnar_texts"]
    fw_y_by_text = layout["fw_y_by_text"]

    result = detect_columns(blocks, page_width)

    # All input blocks must appear in the output (no blocks dropped).
    assert len(result) == len(blocks), (
        f"detect_columns() returned {len(result)} blocks but input had "
        f"{len(blocks)} blocks.\n"
        f"  page_width={page_width}"
    )

    # Build index lookup: block text → position in result list.
    result_index = {b["text"]: i for i, b in enumerate(result)}

    # For each full-width block F, every columnar block C with y_top(C) > y_top(F)
    # must appear AFTER F in the result.
    for fw_text, fw_y in fw_y_by_text.items():
        fw_idx = result_index[fw_text]

        for col_text in columnar_texts:
            # Find the columnar block's top y from the original blocks list.
            col_block = next(b for b in blocks if b["text"] == col_text)
            col_y = col_block["position"][1]

            if col_y > fw_y:
                col_idx = result_index[col_text]
                assert fw_idx < col_idx, (
                    f"Full-width block '{fw_text}' (y={fw_y:.2f}, result_idx={fw_idx}) "
                    f"should precede columnar block '{col_text}' "
                    f"(y={col_y:.2f}, result_idx={col_idx}), but it does not.\n"
                    f"  page_width={page_width:.2f}\n"
                    f"  result order: {[b['text'] for b in result]}"
                )
