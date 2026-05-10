# -*- coding: utf-8 -*-
"""
Shared test fixtures for translation-layout-quality tests.

This module provides:
- A minimal fitz.Page mock for PDF testing
- A sample DOCX Document factory for DOCX testing
- A small block-list factory for general translation testing
"""

import pytest


def pytest_configure(config):
    """Register custom marks to avoid PytestUnknownMarkWarning."""
    config.addinivalue_line(
        "markers",
        "slow: marks tests as slow (e.g. end-to-end pipeline tests that load the NLLB model)",
    )
from unittest.mock import MagicMock, Mock
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


@pytest.fixture
def mock_fitz_page():
    """
    Returns a minimal mock of a fitz.Page object with:
    - get_pixmap() method
    - rect attribute (with width and height)
    - get_text() method returning a dict structure
    """
    page = MagicMock()
    
    # Mock rect attribute
    page.rect = MagicMock()
    page.rect.width = 612.0  # Standard US Letter width in points
    page.rect.height = 792.0  # Standard US Letter height in points
    
    # Mock get_pixmap() method
    pixmap = MagicMock()
    pixmap.width = 612
    pixmap.height = 792
    # Mock pixel() method to return white color by default
    pixmap.pixel = MagicMock(return_value=(255, 255, 255))
    page.get_pixmap = MagicMock(return_value=pixmap)
    
    # Mock get_text() method
    page.get_text = MagicMock(return_value={
        "blocks": [
            {
                "type": 0,  # text block
                "bbox": [50, 50, 500, 100],
                "lines": [
                    {
                        "spans": [
                            {
                                "text": "Sample text",
                                "size": 12.0,
                                "color": 0,
                                "flags": 0,
                                "font": "Helvetica"
                            }
                        ]
                    }
                ]
            }
        ]
    })
    
    return page


@pytest.fixture
def sample_docx_factory():
    """
    Returns a factory function that creates a minimal DOCX Document with:
    - Configurable number of paragraphs
    - Optional table
    - Basic formatting (bold, font size, alignment, spacing)
    
    Usage:
        doc = sample_docx_factory(num_paragraphs=3, include_table=True)
    """
    def _create_docx(num_paragraphs=2, include_table=False, 
                     paragraph_texts=None, table_data=None):
        """
        Create a sample DOCX document.
        
        Args:
            num_paragraphs: Number of paragraphs to create
            include_table: Whether to include a table
            paragraph_texts: Optional list of paragraph text strings
            table_data: Optional 2D list for table content [[row1], [row2], ...]
        
        Returns:
            Document object
        """
        doc = Document()
        
        # Add paragraphs
        if paragraph_texts is None:
            paragraph_texts = [f"Paragraph {i+1} text." for i in range(num_paragraphs)]
        
        for i, text in enumerate(paragraph_texts[:num_paragraphs]):
            para = doc.add_paragraph(text)
            
            # Add some formatting variety
            if i == 0:
                para.style = 'Heading 1'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para.style = 'Normal'
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
            
            # Make first run bold
            if para.runs:
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(12)
                para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Add table if requested
        if include_table:
            if table_data is None:
                table_data = [
                    ["Header 1", "Header 2", "Header 3"],
                    ["Row 1 Col 1", "Row 1 Col 2", "Row 1 Col 3"],
                    ["Row 2 Col 1", "Row 2 Col 2", "Row 2 Col 3"]
                ]
            
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    
                    # Make header row bold
                    if row_idx == 0 and cell.paragraphs:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                            run.font.size = Pt(11)
        
        return doc
    
    return _create_docx


@pytest.fixture
def block_list_factory():
    """
    Returns a factory function that creates a list of block dictionaries
    suitable for testing translation components.
    
    Usage:
        blocks = block_list_factory(count=5, block_type='paragraph')
    """
    def _create_blocks(count=3, block_type='paragraph', include_position=False,
                       include_style=True, texts=None):
        """
        Create a list of block dictionaries.
        
        Args:
            count: Number of blocks to create
            block_type: Type of blocks ('paragraph' or 'table_cell')
            include_position: Whether to include position/page fields (for PDF)
            include_style: Whether to include style metadata
            texts: Optional list of text strings for blocks
        
        Returns:
            List of block dictionaries
        """
        blocks = []
        
        if texts is None:
            texts = [f"This is block number {i+1} with some sample text." 
                    for i in range(count)]
        
        for i, text in enumerate(texts[:count]):
            block = {
                "type": block_type,
                "text": text
            }
            
            if include_position:
                # Create non-overlapping vertical positions
                y_start = 50 + (i * 60)
                block["position"] = [50, y_start, 500, y_start + 50]
                block["page"] = 0
            
            if include_style:
                block["style"] = {
                    "font_size": 12.0,
                    "font": "Helvetica",
                    "color": 0,
                    "bold": False
                }
            
            # Add table-specific fields if needed
            if block_type == 'table_cell':
                block["table_index"] = 0
                block["row"] = i // 3  # Assume 3 columns
                block["col"] = i % 3
                block["row_span"] = 1
                block["col_span"] = 1
            
            blocks.append(block)
        
        return blocks
    
    return _create_blocks
